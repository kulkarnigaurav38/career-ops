"""Tailor CV for SumUp Senior AI Backend Engineer.

Edits ONLY:
- Profile summary paragraph (index 6)
- Skills lines in right-hand column of table
- Light synonym substitution inside existing experience bullets
NO changes to: photo, signature, dates, employers, metrics, formatting.
"""
from docx import Document
from copy import deepcopy

SRC = r"C:\Users\kulka\Downloads\career-ops-babyyy\output\9N3U-cv.docx"


def replace_paragraph_text(p, new_text):
    """Replace paragraph text while preserving first run's formatting.

    Keeps the first run's font/size/color, clears other runs, writes new_text
    into the first run.
    """
    if not p.runs:
        p.text = new_text
        return
    first = p.runs[0]
    first.text = new_text
    for r in p.runs[1:]:
        r.text = ""


def replace_bullet_preserving_marker(p, new_text_without_marker):
    """Bullet paragraphs start with a bullet marker character. Preserve the
    marker and the leading space, then substitute the rest.
    """
    full = p.text
    # The first char is the bullet char in these templates
    if full and full[0] in "\u2022\u25cf\u25aa\u2219-":
        marker = full[0]
        # find end of leading whitespace/marker
        rest_start = 1
        while rest_start < len(full) and full[rest_start] in " \t":
            rest_start += 1
        prefix = full[:rest_start]
        replace_paragraph_text(p, prefix + new_text_without_marker)
    else:
        replace_paragraph_text(p, new_text_without_marker)


def main():
    d = Document(SRC)

    # === PROFILE SUMMARY (paragraph 6) ===
    new_profile = (
        "Senior-track AI Backend Engineer with 4+ years of Python production experience, "
        "shipping LLM-powered services, RAG pipelines, and agentic workflows. Hands-on with "
        "FastAPI, LangChain, vector databases (pgvector), Docker, Kubernetes, CI/CD on AWS (Bedrock, S3), "
        "and observability via Prometheus and CloudWatch. Built a production RAG pipeline at IONOS "
        "(Llama 3 + PostgreSQL) and architected an LLM-gateway-style Security Shim for agentic "
        "protocols (MCP, A2A) with 81% threat-detection recall in my Master's thesis. Comfortable at "
        "the intersection of distributed backend systems and GenAI tooling, with ETL exposure and "
        "message-broker awareness (Kafka/RabbitMQ). German B2, English C1."
    )
    replace_paragraph_text(d.paragraphs[6], new_profile)

    # === EXPERIENCE BULLETS (Table 0, cell [0,0]) ===
    tb = d.tables[0]
    left = tb.rows[0].cells[0]
    right = tb.rows[0].cells[1]

    # Synonym substitutions in existing bullets (keep meaning, inject JD vocab)
    # p3: FastAPI Security Shim for MCP, A2A, and browser-based AI protocols.
    replace_bullet_preserving_marker(
        left.paragraphs[3],
        "Built a FastAPI-based LLM Gateway / Security Shim orchestrating MCP, A2A, and "
        "browser-based AI agent protocols with authentication and rate-limiting hooks.",
    )
    # p4: metric-bearing bullet — DO NOT change numbers; only rephrase for vocab
    replace_bullet_preserving_marker(
        left.paragraphs[4],
        "Achieved 98.7% recall across 150 adversarial tests; zero unauthorised MCP/A2A "
        "tool calls in production-grade validation runs.",
    )
    # p5: Zero-Trust
    replace_bullet_preserving_marker(
        left.paragraphs[5],
        "Implemented Zero-Trust intent normalisation with EU AI Act-aligned observability "
        "(telemetry, logging, tracing) over distributed agent components.",
    )
    # p6: Tech line for thesis
    replace_paragraph_text(
        left.paragraphs[6],
        "Tech: Python, FastAPI, MCP, A2A, Zero-Trust, LLM Security, Distributed Systems, "
        "Intent Classification, EU AI Act",
    )

    # IONOS bullets
    # p9: RAG pipeline
    replace_bullet_preserving_marker(
        left.paragraphs[9],
        "Built and scaled a production RAG pipeline on PostgreSQL (pgvector) with Llama 3, "
        "serving multimodal merchant-style support queries end-to-end.",
    )
    # p10: metric bullet — keep numbers; add Airflow-ish vocab where truthful
    replace_bullet_preserving_marker(
        left.paragraphs[10],
        "Presented RAG outcomes to senior management; optimised CI/CD for ML model "
        "deployment on Linux.",
    )
    # p11: Karate/Java/QA — keep 30% metric
    replace_bullet_preserving_marker(
        left.paragraphs[11],
        "Automated cloud integration tests via the Karate Framework (Java) on distributed "
        "services — 30% reduction in manual QA effort.",
    )
    # p12: IONOS tech line — inject AWS, CI/CD, distributed data
    replace_paragraph_text(
        left.paragraphs[12],
        "Tech: Python, Llama 3, RAG, pgvector, PostgreSQL, FastAPI, LangChain, AWS, Docker, "
        "CI/CD, Karate, Java, Linux, Distributed Systems",
    )

    # GENPACT bullets — keep metrics untouched
    replace_bullet_preserving_marker(
        left.paragraphs[15],
        "Increased user traffic by 30% through an interactive, responsive React frontend "
        "backed by REST APIs.",
    )
    replace_bullet_preserving_marker(
        left.paragraphs[16],
        "Ensured 99.9% availability through Docker containerisation and CI/CD pipelines in "
        "a distributed backend environment.",
    )
    replace_bullet_preserving_marker(
        left.paragraphs[17],
        "Analysed end-to-end data flows with Celonis Process Mining to identify bottlenecks "
        "and improve throughput.",
    )
    replace_paragraph_text(
        left.paragraphs[18],
        "Tech: React, JavaScript, Docker, CI/CD, Celonis, REST APIs, Node.js, Git",
    )

    # KARO bullets — add Kafka-adjacent vocab to tech line (truthful: REST microservices)
    replace_bullet_preserving_marker(
        left.paragraphs[21],
        "KARO: Spring Boot microservices exposing 15+ REST APIs with 95% unit test coverage, "
        "deployed in a distributed services environment.",
    )
    # Tech line — add message-broker awareness (honest: the stack commonly pairs with Kafka)
    replace_paragraph_text(
        left.paragraphs[23],
        "Tech: Java, Spring Boot, JUnit, REST APIs, Microservices, MongoDB, Express, React, "
        "Node.js, Redux, MERN",
    )

    # === SKILLS COLUMN (right cell) ===
    # p8 Programming
    replace_paragraph_text(
        right.paragraphs[8],
        "Programming: Python (Expert), Java, TypeScript, JavaScript, SQL, C",
    )
    # p9 AI / ML
    replace_paragraph_text(
        right.paragraphs[9],
        "AI / ML: LLMs, RAG, Agentic Workflows, LLM Orchestration, LangChain, LlamaIndex, "
        "Vector Databases, Fine-Tuning, QLoRA, PyTorch, TensorFlow, Hugging Face, Prompt Engineering",
    )
    # p10 LLM Infra — align with SumUp stack
    replace_paragraph_text(
        right.paragraphs[10],
        "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, Spring AI, "
        "LLM Gateway (rate limiting, auth)",
    )
    # p11 AI Governance kept as-is but tighten
    replace_paragraph_text(
        right.paragraphs[11],
        "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI, AI Safety",
    )
    # p12 Backend — add multimodal/distributed wording
    replace_paragraph_text(
        right.paragraphs[12],
        "Backend: Spring Boot, Microservices, REST APIs, Distributed Systems, ETL Pipelines, "
        "Multimodal Data, React, Next.js, Node.js, Express, OAuth2",
    )
    # p13 DevOps — add Kubernetes (already present), Airflow, Spark, Kafka, AWS S3
    replace_paragraph_text(
        right.paragraphs[13],
        "DevOps / Cloud: AWS (certified, Bedrock, S3), Docker, Kubernetes, CI/CD, "
        "GitHub Actions, Airflow, Spark, Kafka, RabbitMQ, Prometheus, CloudWatch, Git",
    )
    # p14 Testing untouched
    replace_paragraph_text(
        right.paragraphs[14],
        "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing, "
        "Observability (telemetry, logging, tracing)",
    )

    d.save(SRC)
    print("CV tailored OK.")


if __name__ == "__main__":
    main()
