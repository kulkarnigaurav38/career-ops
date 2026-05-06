"""
Tailor helper for batch CV + cover letter generation.

Usage:
    python tailor_helper.py cv <master_docx> <out_docx> <config_json>
    python tailor_helper.py cl <master_docx> <out_docx> <config_json>

config_json is a path to a JSON file with keys:
  For CV:
    - "profile": full new Profile Summary paragraph text
    - "bullet_replacements": dict of original_text_fragment -> new_text
        (we match by paragraph index range 3..20 inclusive — NOT 21/22 which have hyperlinks)
    - "skills": dict of skill_category_line_prefix -> full new line
        (replaces skill lines in right column, e.g. "Programming:", "AI / ML:")
  For CL:
    - "company_name": str
    - "company_address_lines": list[str] (1-3 lines)
    - "date_line": str (e.g. "Leonberg, April 21, 2026")
    - "subject": str (replaces "Application for ..." / "Bewerbung ...")
    - "greeting": str (optional, e.g. "Dear Hiring Team,")
    - "body_paragraphs": list[str] (6 paragraphs — hook, 3 fit paragraphs, close, sign-off line)
        Length must match the 6 body paragraphs in the master (indices 7..12)
    - "closing_paragraph": str (paragraph 13 — "My signature follows..." line)
"""

import json
import sys
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn


def set_paragraph_text_preserve_style(paragraph, new_text):
    """Replace paragraph text, preserving the style of the first run.
    Keeps bold/italic/font of run[0]; deletes all other runs.
    Does NOT touch hyperlinks."""
    # Check if paragraph has a hyperlink — if yes, skip (caller's problem)
    if paragraph._element.findall(qn('w:hyperlink')):
        raise ValueError(f"Paragraph contains hyperlink, refusing to edit: {paragraph.text[:80]}")

    runs = paragraph.runs
    if not runs:
        # empty paragraph — just add a run
        paragraph.add_run(new_text)
        return

    # Keep first run, clear its text, set new text
    runs[0].text = new_text
    # Remove all other runs
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def tailor_cv(master_path, out_path, config_path):
    with open(config_path, 'r', encoding='utf-8') as f:
        cfg = json.load(f)

    # Copy master to out
    import shutil
    shutil.copy2(master_path, out_path)

    doc = Document(out_path)

    # 1. Profile Summary — paragraph 6 (top-level)
    if 'profile' in cfg:
        new_profile = cfg['profile']
        p = doc.paragraphs[6]
        set_paragraph_text_preserve_style(p, new_profile)

    # 2. Experience bullets in table 0, cell [0,0], paragraphs 3..20 (SKIP 21 & 22)
    # Also skip Tech: lines if caller didn't include them
    if 'bullet_replacements' in cfg:
        t = doc.tables[0]
        cell = t.rows[0].cells[0]
        # Safe indices to edit: 3-5 (thesis), 9-11 (IONOS), 15-17 (GENPACT)
        # Tech lines: 6, 12, 18, 23 — also editable (no hyperlinks)
        # PROJECTS bullets: 27-29, 32-33, 36-37, 40-41, 44-45, 48 (+ Stack lines)
        # DO NOT touch 21, 22 (KARO/JITSIE with hyperlinks)
        safe_indices = list(range(3, 21)) + list(range(23, 50))
        # Remove 21, 22 explicitly
        safe_indices = [i for i in safe_indices if i not in (21, 22)]

        for idx in safe_indices:
            if idx >= len(cell.paragraphs):
                continue
            p = cell.paragraphs[idx]
            current = p.text
            # apply replacements (substring match, case-sensitive)
            new_text = current
            matched = False
            for orig_frag, new_frag in cfg['bullet_replacements'].items():
                if orig_frag in new_text:
                    new_text = new_text.replace(orig_frag, new_frag)
                    matched = True
            if matched and new_text != current:
                try:
                    set_paragraph_text_preserve_style(p, new_text)
                except ValueError as e:
                    print(f"  skip p{idx}: {e}", file=sys.stderr)

    # 3. Skills — table 0, cell [0,1], paragraphs 8..14 (lines starting with Programming:, AI/ML:, etc.)
    if 'skills' in cfg:
        t = doc.tables[0]
        cell = t.rows[0].cells[1]
        for idx in range(8, 15):
            if idx >= len(cell.paragraphs):
                continue
            p = cell.paragraphs[idx]
            current = p.text
            for prefix, new_line in cfg['skills'].items():
                if current.startswith(prefix):
                    try:
                        set_paragraph_text_preserve_style(p, new_line)
                    except ValueError as e:
                        print(f"  skip skill p{idx}: {e}", file=sys.stderr)
                    break

    doc.save(out_path)
    print(f"CV tailored: {out_path}")


def tailor_cl(master_path, out_path, config_path):
    with open(config_path, 'r', encoding='utf-8') as f:
        cfg = json.load(f)

    import shutil
    shutil.copy2(master_path, out_path)

    doc = Document(out_path)

    # Company address — table 0, cell[0,0] paragraphs 0..2
    # Date line — table 0, cell[0,1] paragraph 0
    t = doc.tables[0]
    left = t.rows[0].cells[0]
    right = t.rows[0].cells[1]

    # Company name + address
    if 'company_address_lines' in cfg:
        addr_lines = cfg['company_address_lines']
        for i, line in enumerate(addr_lines[:3]):
            if i < len(left.paragraphs):
                try:
                    set_paragraph_text_preserve_style(left.paragraphs[i], line)
                except ValueError as e:
                    print(f"  skip addr p{i}: {e}", file=sys.stderr)
        # If fewer lines provided than exist, blank out extras
        for i in range(len(addr_lines), 3):
            if i < len(left.paragraphs):
                try:
                    set_paragraph_text_preserve_style(left.paragraphs[i], "")
                except ValueError:
                    pass

    # Date line
    if 'date_line' in cfg and len(right.paragraphs) > 0:
        try:
            set_paragraph_text_preserve_style(right.paragraphs[0], cfg['date_line'])
        except ValueError as e:
            print(f"  skip date: {e}", file=sys.stderr)

    # Main body — top-level paragraphs
    # 5: subject, 6: greeting, 7-12: body (6 paragraphs), 13: closing line, 14: "Kind regards,"
    paragraphs = doc.paragraphs

    if 'subject' in cfg and len(paragraphs) > 5:
        try:
            set_paragraph_text_preserve_style(paragraphs[5], cfg['subject'])
        except ValueError as e:
            print(f"  skip subject: {e}", file=sys.stderr)

    if 'greeting' in cfg and len(paragraphs) > 6:
        try:
            set_paragraph_text_preserve_style(paragraphs[6], cfg['greeting'])
        except ValueError as e:
            print(f"  skip greeting: {e}", file=sys.stderr)

    if 'body_paragraphs' in cfg:
        body = cfg['body_paragraphs']
        # Master has 6 body paragraphs at indices 7..12
        for i, text in enumerate(body[:6]):
            idx = 7 + i
            if idx < len(paragraphs):
                try:
                    set_paragraph_text_preserve_style(paragraphs[idx], text)
                except ValueError as e:
                    print(f"  skip body p{idx}: {e}", file=sys.stderr)

    if 'closing_paragraph' in cfg and len(paragraphs) > 13:
        try:
            set_paragraph_text_preserve_style(paragraphs[13], cfg['closing_paragraph'])
        except ValueError as e:
            print(f"  skip closing: {e}", file=sys.stderr)

    doc.save(out_path)
    print(f"Cover letter tailored: {out_path}")


if __name__ == '__main__':
    mode = sys.argv[1]
    master = sys.argv[2]
    out = sys.argv[3]
    cfg = sys.argv[4]
    if mode == 'cv':
        tailor_cv(master, out, cfg)
    elif mode == 'cl':
        tailor_cl(master, out, cfg)
    else:
        print(f"Unknown mode: {mode}", file=sys.stderr)
        sys.exit(1)
