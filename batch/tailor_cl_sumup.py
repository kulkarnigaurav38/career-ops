"""Tailor Cover Letter for SumUp Senior AI Backend Engineer."""
from docx import Document

SRC = r"C:\Users\kulka\Downloads\career-ops-babyyy\output\9N3U-coverletter.docx"


def replace_paragraph_text(p, new_text):
    if not p.runs:
        p.text = new_text
        return
    first = p.runs[0]
    first.text = new_text
    for r in p.runs[1:]:
        r.text = ""


def main():
    d = Document(SRC)

    # === HEADER TABLE (addressee + date) ===
    tb = d.tables[0]
    addr_cell = tb.rows[0].cells[0]
    date_cell = tb.rows[0].cells[1]

    replace_paragraph_text(addr_cell.paragraphs[0], "SumUp")
    replace_paragraph_text(addr_cell.paragraphs[1], "Talent Acquisition Team")
    replace_paragraph_text(addr_cell.paragraphs[2], "Berlin, Germany")

    replace_paragraph_text(date_cell.paragraphs[0], "Leonberg, April 21, 2026")

    # === SUBJECT LINE (paragraph 5) ===
    replace_paragraph_text(
        d.paragraphs[5],
        "Application for Senior AI Backend Engineer (f/m/d) — SumUp Edge",
    )

    # Greeting paragraph 6 stays the same: "Dear Sir or Madam,"

    # === HOOK PARAGRAPH (7) ===
    replace_paragraph_text(
        d.paragraphs[7],
        "SumUp Edge caught my attention because it matches exactly what I've been building "
        "for the past two years: AI systems that turn raw data into proactive, real-time "
        "decisions. At IONOS I shipped a production RAG pipeline, and my Master's thesis "
        "is a working LLM-gateway / Security Shim orchestrating MCP and A2A agents — "
        "the same class of backend problem SumUp Edge is solving for merchants.",
    )

    # === FIT INTRO (8) ===
    replace_paragraph_text(
        d.paragraphs[8],
        "Rather than claim I'm the right fit, here is how my experience maps directly to "
        "the Senior AI Backend Engineer role:",
    )

    # === BULLETS (9, 10, 11) ===
    # Bullet 9 — LLM Gateway + Orchestration + Vector DB
    replace_paragraph_text(
        d.paragraphs[9],
        "You are building an orchestrator service, LLM Gateway, and vector database "
        "infrastructure for SumUp Edge. At IONOS I designed and implemented a RAG "
        "pipeline with Llama 3 on PostgreSQL (pgvector). My Master's thesis extends "
        "that into a FastAPI-based Security Shim that coordinates agentic protocols "
        "(MCP, A2A, DOM) with authentication, rate limiting, and 81% threat-detection "
        "recall — exactly the orchestration + gateway concerns described in the JD.",
    )
    # Bullet 10 — Python + scalable backend + ETL + multimodal APIs
    replace_paragraph_text(
        d.paragraphs[10],
        "You need advanced Python with production backend scaling, ETL pipelines, and "
        "high-performance APIs for multimodal data. I write Python every day — FastAPI "
        "inference APIs for my fine-tuned Llama 3 DCF engine, ETL for the Dibuco news "
        "intelligence platform (scraping, clustering, transformer sentiment), and "
        "distributed REST services at IONOS and KARO. I care about clean, testable, "
        "observable code.",
    )
    # Bullet 11 — Kubernetes, Docker, Airflow/Spark, AWS, CI/CD, monitoring
    replace_paragraph_text(
        d.paragraphs[11],
        "You run on Kubernetes, Docker, Airflow, Spark, and AWS, with Prometheus and "
        "CloudWatch for observability. I'm AWS-certified (Bedrock), have shipped "
        "containerised services at GENPACT with 99.9% availability, optimised CI/CD "
        "for ML models at IONOS, and built my thesis around telemetry, logging, and "
        "tracing for agent systems. Kafka/RabbitMQ sit naturally on the same mental "
        "model I already use for distributed data flow.",
    )

    # === CLOSE PARAGRAPH (12) ===
    replace_paragraph_text(
        d.paragraphs[12],
        "SumUp Edge is the kind of product I'd want to be in the trenches on — real "
        "merchants, real data, real latency constraints, and an AI layer that has to "
        "actually work in production. I'd bring immediately productive AI and backend "
        "expertise, a builder mentality shaped by 4 years across full-stack, cloud, "
        "and GenAI, and the rigour that comes from researching AI agent security for "
        "my thesis. Berlin works for me and I'm happy with the office-first setup.",
    )

    # Paragraph 13 (signature lead-in), 14 (Kind regards), 15 (name) stay untouched.

    d.save(SRC)
    print("Cover letter tailored OK.")


if __name__ == "__main__":
    main()
