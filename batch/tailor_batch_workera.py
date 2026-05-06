"""Worker A batch tailoring: tailor CV + cover letter from DOCX masters.

Usage:
    python tailor_batch_workera.py <job_key>

Each job_key loads a dict of tailor config below, copies the master, edits
Profile, bullets 1-20 (skips 21/22 with hyperlinks), Skills, then writes cover
letter body. Saves DOCX to output/{REF}-cv.docx and output/{REF}-coverletter.docx.

NO formatting, photo, signature, dates, employers, metrics changes.
"""
from docx import Document
import shutil
import sys
import os

ROOT = r"C:\Users\kulka\Downloads\career-ops-babyyy"
CV_EN_MASTER = os.path.join(ROOT, "templates", "cv", "CV_Gaurav_Kulkarni_EN.docx")
CL_EN_MASTER = os.path.join(ROOT, "templates", "cv", "CoverLetter_Gaurav_Kulkarni_DeutscheBoerse_EN.docx")
CV_DE_MASTER = os.path.join(ROOT, "templates", "cv", "Lebenslauf_Gaurav_Kulkarni_DE.docx")
CL_DE_MASTER = os.path.join(ROOT, "templates", "cv", "Anschreiben_Gaurav_Kulkarni_DeutscheBoerse_DE.docx")


def replace_paragraph_text(p, new_text):
    """Replace paragraph text while preserving first run's formatting."""
    if not p.runs:
        p.text = new_text
        return
    first = p.runs[0]
    first.text = new_text
    for r in p.runs[1:]:
        r.text = ""


def replace_bullet_preserving_marker(p, new_text_without_marker):
    """Preserve leading bullet character + whitespace."""
    full = p.text
    if full and full[0] in "\u2022\u25cf\u25aa\u2219-":
        rest_start = 1
        while rest_start < len(full) and full[rest_start] in " \t":
            rest_start += 1
        prefix = full[:rest_start]
        replace_paragraph_text(p, prefix + new_text_without_marker)
    else:
        replace_paragraph_text(p, new_text_without_marker)


def tailor_cv_en(ref, cfg):
    out = os.path.join(ROOT, "output", f"{ref}-cv.docx")
    shutil.copy2(CV_EN_MASTER, out)
    d = Document(out)

    # Profile (P6)
    replace_paragraph_text(d.paragraphs[6], cfg["profile"])

    tb = d.tables[0]
    left = tb.rows[0].cells[0]
    right = tb.rows[0].cells[1]

    # Thesis bullets P3-P5
    b = cfg.get("thesis_bullets", [])
    for i, text in enumerate(b):
        replace_bullet_preserving_marker(left.paragraphs[3 + i], text)
    if cfg.get("thesis_tech"):
        replace_paragraph_text(left.paragraphs[6], cfg["thesis_tech"])

    # IONOS bullets P9-P11
    b = cfg.get("ionos_bullets", [])
    for i, text in enumerate(b):
        replace_bullet_preserving_marker(left.paragraphs[9 + i], text)
    if cfg.get("ionos_tech"):
        replace_paragraph_text(left.paragraphs[12], cfg["ionos_tech"])

    # GENPACT bullets P15-P17
    b = cfg.get("genpact_bullets", [])
    for i, text in enumerate(b):
        replace_bullet_preserving_marker(left.paragraphs[15 + i], text)
    if cfg.get("genpact_tech"):
        replace_paragraph_text(left.paragraphs[18], cfg["genpact_tech"])

    # KARO/JITSIE (P21, P22) — DO NOT EDIT (hyperlinks)
    # Tech line P23 is editable
    if cfg.get("iit_tech"):
        replace_paragraph_text(left.paragraphs[23], cfg["iit_tech"])

    # Skills — right cell
    if cfg.get("skills"):
        s = cfg["skills"]
        # P8 Programming
        if "programming" in s:
            replace_paragraph_text(right.paragraphs[8], s["programming"])
        if "ai_ml" in s:
            replace_paragraph_text(right.paragraphs[9], s["ai_ml"])
        if "llm_infra" in s:
            replace_paragraph_text(right.paragraphs[10], s["llm_infra"])
        if "ai_gov" in s:
            replace_paragraph_text(right.paragraphs[11], s["ai_gov"])
        if "backend" in s:
            replace_paragraph_text(right.paragraphs[12], s["backend"])
        if "devops" in s:
            replace_paragraph_text(right.paragraphs[13], s["devops"])
        if "testing" in s:
            replace_paragraph_text(right.paragraphs[14], s["testing"])

    d.save(out)
    return out


def tailor_cl_en(ref, cfg):
    out = os.path.join(ROOT, "output", f"{ref}-coverletter.docx")
    shutil.copy2(CL_EN_MASTER, out)
    d = Document(out)

    # P5: subject
    if cfg.get("subject"):
        replace_paragraph_text(d.paragraphs[5], cfg["subject"])

    # P6: greeting (keep default "Dear Sir or Madam," unless overridden)
    if cfg.get("greeting"):
        replace_paragraph_text(d.paragraphs[6], cfg["greeting"])

    # P7-P12: body (6 paragraphs)
    body = cfg.get("body", [])
    for i, text in enumerate(body[:6]):
        replace_paragraph_text(d.paragraphs[7 + i], text)

    # P13: closing is standard, leave unless overridden
    if cfg.get("closing"):
        replace_paragraph_text(d.paragraphs[13], cfg["closing"])

    d.save(out)
    return out


# ============================================================================
# JOB CONFIGS — one dict per REF
# ============================================================================

JOBS = {}

# ---------------- 009 Sokra — Founding Engineer (Full-Stack AI/ML) ----------------
JOBS["0UY1"] = {
    "lang": "en",
    "profile": (
        "Founding-engineer-ready full-stack AI/ML builder with 4+ years shipping production software, "
        "now focused on LLM agents, RAG pipelines, and adaptive user experiences. Hands-on with TypeScript, "
        "React, Next.js, Python, FastAPI, WebSockets, and end-to-end 0-to-1 delivery from prototype to "
        "deployed product. Built a production RAG pipeline at IONOS (Llama 3 + PostgreSQL), a Security "
        "Shim for agentic protocols (MCP, A2A) with 81% threat-detection recall, and a full Next.js + "
        "Prisma parking platform for an airport client. Comfortable owning architecture, iterating with "
        "real users (50+ schools class of problem), and moving fast. CS background, German B2, ready to "
        "relocate to Berlin."
    ),
    "thesis_bullets": [
        "Built a FastAPI-based LLM agent gateway / Security Shim orchestrating MCP, A2A, and browser-based AI protocols with authentication and policy hooks.",
        "Achieved 98.7% recall across 150 adversarial tests; zero unauthorised MCP/A2A tool calls in end-to-end validation runs.",
        "Implemented Zero-Trust intent normalisation with EU AI Act-aligned observability across distributed agent components.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, MCP, A2A, LLM Agents, Zero-Trust, Intent Classification, EU AI Act",
    "ionos_bullets": [
        "Built a production RAG pipeline on PostgreSQL (pgvector) with Llama 3 powering adaptive, context-aware responses.",
        "Presented RAG outcomes to senior management; optimised CI/CD for ML model deployment.",
        "Automated cloud integration tests via the Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, Llama 3, RAG, pgvector, PostgreSQL, FastAPI, LangChain, LLM Agents, Adaptive Learning, CI/CD, AWS",
    "genpact_bullets": [
        "Shipped 0-to-1 interactive, responsive React frontend — 30% traffic increase end-to-end.",
        "Ensured 99.9% availability through Docker and CI/CD pipelines in a fast-iteration environment.",
        "Analysed user data flows with Celonis Process Mining to identify product bottlenecks.",
    ],
    "genpact_tech": "Tech: React, TypeScript, Next.js, JavaScript, Docker, CI/CD, Celonis, REST APIs, Node.js, Git",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, MongoDB, Express, React, Node.js, Redux, MERN, Full-Stack",
    "skills": {
        "programming": "Programming: TypeScript, Python (Expert), JavaScript, SQL, Java, C",
        "ai_ml": "AI / ML: LLMs, RAG, LLM Agents, Agentic Workflows, LangChain, Vector Databases, Fine-Tuning, QLoRA, PyTorch, Hugging Face, Adaptive Learning, Prompt Engineering",
        "llm_infra": "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, Spring AI, LLM Gateway",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI, AI Safety",
        "backend": "Full-Stack: Next.js 16, React 19, Node.js, Express, WebSockets, REST APIs, OAuth2, JWT, Prisma, Stripe, Tailwind",
        "devops": "DevOps / Cloud: AWS (certified), Docker, Kubernetes, CI/CD, GitHub Actions, Vercel, Neon, Git, Playwright",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing",
    },
    "subject": "Application for Founding Engineer (Full-Stack · AI/ML) — Sokra",
    "body": [
        "Sokra caught my eye because it's the exact shape of product I've been building for the past two years: AI-native software that adapts to real users — here, students and teachers across 50+ schools. At IONOS I shipped a production RAG pipeline with Llama 3, my Master's thesis is a working FastAPI LLM-agent gateway (MCP, A2A), and I recently delivered a full Next.js + Prisma airport parking platform 0-to-1 for a freelance client. Founding-engineer territory is where I'm most at home.",
        "Here's how my experience maps directly to what Sokra needs in a founding engineer:",
        "You're building an AI-powered adaptive learning engine with LLM agents and RAG. At IONOS I designed and shipped a RAG pipeline with Llama 3 on PostgreSQL (pgvector); my thesis extends that into agent orchestration (MCP, A2A, browser-based AI) with authentication, policy routing, and 81% threat-detection recall. Adaptive content delivery sits naturally on top of that stack.",
        "You need strong TypeScript, React, Next.js, and Python. Python is my daily driver for AI work; my recent airport platform is Next.js 16 + React 19 + TypeScript + Prisma with 22 rate-limited APIs, Stripe payments, Playwright E2E, and WebSockets-class real-time flows. I care about shipping clean, observable code and iterating quickly on real feedback.",
        "You want 0-to-1 builder mentality. GENPACT trained me on high-availability (99.9%) containerised systems; KARO demanded 15+ REST APIs with 95% test coverage on critical deadlines; the LIT Airport client needed the whole stack delivered in four weeks. I ship, I iterate, I talk to users — and I want a team where those reflexes compound.",
        "Sokra's mission — software that teachers and students actually use every day — is the kind of work I'd be proud to put my name on. I'm already based in Germany, B2 German, and ready to be in Berlin. I'd bring immediately productive AI + full-stack engineering, a builder mentality shaped by 4 years across cloud, fullstack, and GenAI, and the ownership instinct that early-stage teams need.",
    ],
}

# ---------------- 010 Jamie — AI Engineer ----------------
JOBS["PYEN"] = {
    "lang": "en",
    "profile": (
        "Product-minded AI Engineer with 4+ years shipping software and 2+ years optimizing AI/ML pipelines "
        "end-to-end. Hands-on with LLMs, prompt engineering, MLOps, RAG, fine-tuning (QLoRA), and production "
        "inference APIs. Built a RAG pipeline at IONOS (Llama 3 + PostgreSQL), a Security Shim for AI agents "
        "with 81% threat-detection recall for my Master's thesis, and a fine-tuned Llama 3 DCF engine served "
        "via FastAPI. Comfortable iterating fast on user feedback, owning both code and prompts, and caring "
        "about product metrics. German B2, English C1, based in Germany and happy to be in Berlin Mitte."
    ),
    "thesis_bullets": [
        "Built a FastAPI Security Shim / LLM pipeline orchestrating MCP, A2A, and browser-based AI agent protocols.",
        "Achieved 98.7% recall across 150 adversarial tests; tuned prompts and policies for reliable, measurable AI behaviour.",
        "Implemented Zero-Trust intent normalisation with observability-first telemetry for AI agent runs.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, MCP, A2A, Prompt Engineering, LLM Evaluation, MLOps, Zero-Trust",
    "ionos_bullets": [
        "Built a production RAG pipeline on PostgreSQL (pgvector) with Llama 3 powering meeting-/note-style contextual answers.",
        "Iterated on prompt design and inference quality based on internal user feedback; optimised CI/CD for ML deployment.",
        "Automated cloud testing via the Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, Llama 3, RAG, pgvector, PostgreSQL, FastAPI, LangChain, Prompt Engineering, MLOps, CI/CD, AWS",
    "genpact_bullets": [
        "Increased user traffic by 30% via a product-focused, iterative React frontend.",
        "Ensured 99.9% availability through Docker containerisation and CI/CD pipelines.",
        "Used Celonis Process Mining to turn user-behaviour data into product improvements.",
    ],
    "genpact_tech": "Tech: React, JavaScript, Docker, CI/CD, Celonis, REST APIs, Product Analytics, Node.js, Git",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, Microservices, MongoDB, Express, React, Node.js, Redux, MERN",
    "skills": {
        "programming": "Programming: Python (Expert, AI/ML), TypeScript, JavaScript, Java, SQL, C",
        "ai_ml": "AI / ML: LLMs, Prompt Engineering, RAG, Agentic Workflows, LangChain, LlamaIndex, Vector Databases, Fine-Tuning, QLoRA, PEFT, PyTorch, Hugging Face, LLM Evaluation, MLOps",
        "llm_infra": "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, Spring AI, Inference APIs",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI, AI Safety",
        "backend": "Backend: Next.js, React, Node.js, Express, REST APIs, OAuth2, Product Analytics, Stripe, Tailwind",
        "devops": "DevOps / Cloud: AWS (certified), Docker, Kubernetes, CI/CD, GitHub Actions, Vercel, Git",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Observability, Telemetry",
    },
    "subject": "Application for AI Engineer — Jamie",
    "body": [
        "Jamie's product — an AI meeting assistant that has to actually work for real users — is exactly the kind of AI I want to be building: user-facing, feedback-driven, and judged by product metrics rather than benchmarks. At IONOS I shipped a production RAG pipeline on Llama 3, and my Master's thesis is a working LLM-agent Security Shim with MCP and A2A orchestration. The muscle Jamie needs — turning requirements into code and prompts that hold up in production — is the muscle I've been training.",
        "Here is how my experience maps to the AI Engineer role:",
        "You need an engineer who converts user requirements into robust code and prompts for a meeting note platform. At IONOS I translated real support scenarios into a RAG pipeline (Llama 3 + pgvector) that surfaced contextual answers. My thesis work is pure prompt + policy engineering — orchestrating agentic protocols and measuring with 150 adversarial tests, 98.7% recall. I think about prompts and code together, not separately.",
        "You want 2+ years of software engineering with shipped projects. I have 4 years across GenAI, cloud, and full-stack: IONOS (RAG in production), KARO (15+ REST APIs, 95% test coverage), GENPACT (30% traffic lift, 99.9% uptime), a fine-tuned Llama 3 DCF engine served via FastAPI, and a freelance airport booking platform in Next.js 16 + TypeScript + Prisma + Stripe. MLOps and prompt-eng are part of my daily loop.",
        "You move fast and iterate on user feedback. I'm at my sharpest when I can sit close to users — KARO had me talking directly to the founders on deadline; the airport client had me iterating with the ops team weekly. Berlin Mitte, multiple days a week, is a feature for me, not a bug.",
        "Jamie looks like a team I'd want to be at in year-one of a product that could genuinely compound. I'd bring immediately productive AI + full-stack engineering, a builder mentality shaped across 4 years of shipping, and a real appetite to own features end-to-end — code, prompts, metrics.",
    ],
}

# ---------------- 011 Machine Learning Reply — GenAI Engineer (m/w/d) ----------------
JOBS["WF4I"] = {
    "lang": "en",
    "profile": (
        "GenAI Engineer with 4+ years of software experience and hands-on LLM, NLP, and multi-agent "
        "production work. Strong in Python, prompt engineering, fine-tuning (QLoRA), RAG, and cloud "
        "deployment on AWS. Built a production RAG pipeline at IONOS (Llama 3 + PostgreSQL), a FastAPI "
        "LLM-agent Security Shim (MCP, A2A) for my Master's thesis with 81% threat-detection recall, and "
        "a fine-tuned Llama 3 DCF reasoning engine. Comfortable translating client business requirements "
        "into AI solutions, presenting to senior stakeholders, and mentoring juniors. German B2, English C1, "
        "happy to be in Munich."
    ),
    "thesis_bullets": [
        "Built a FastAPI LLM-agent Security Shim orchestrating MCP, A2A, and browser-based multi-agent protocols.",
        "Achieved 98.7% recall across 150 adversarial tests; zero unauthorised MCP/A2A tool calls in validation.",
        "Implemented Zero-Trust intent normalisation with EU AI Act-aligned observability for GenAI systems.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, MCP, A2A, LLM Agents, Multi-Agent Systems, Zero-Trust, EU AI Act",
    "ionos_bullets": [
        "Designed and shipped a production RAG pipeline on PostgreSQL (pgvector) with Llama 3 for internal enterprise support.",
        "Presented RAG outcomes to senior management — translated technical results into business impact.",
        "Automated cloud integration tests via Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, Llama 3, RAG, pgvector, PostgreSQL, FastAPI, LangChain, Hugging Face, AWS, Docker, CI/CD",
    "genpact_bullets": [
        "Increased user traffic by 30% through an interactive, responsive React frontend for an HDAX-class enterprise.",
        "Ensured 99.9% availability through Docker containerisation and CI/CD.",
        "Analysed end-to-end data flows with Celonis Process Mining to accelerate enterprise client workflows.",
    ],
    "genpact_tech": "Tech: React, JavaScript, Docker, CI/CD, Celonis, REST APIs, Node.js, Git, Enterprise Workflows",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, Microservices, MongoDB, Express, React, Node.js, Redux, MERN",
    "skills": {
        "programming": "Programming: Python (Expert), SQL, PySpark, TypeScript, JavaScript, Java, C",
        "ai_ml": "GenAI / LLMs: LLMs, Multi-Agent LLM Applications, RAG, Agentic Workflows, Prompt Engineering, Context Engineering, Fine-Tuning, QLoRA, PEFT, NLP, LangChain, Hugging Face, OpenAI APIs, LlamaIndex, PyTorch, Transformers, Knowledge Graphs",
        "llm_infra": "LLM Infra: AWS Bedrock, Azure OpenAI, GCP, PostgreSQL, pgvector, FastAPI, Llama 3, Databricks",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI, AI Safety",
        "backend": "Backend: Spring Boot, Microservices, REST APIs, Next.js, React, Node.js, Express, OAuth2",
        "devops": "Cloud / DevOps: AWS (certified, Bedrock), Azure, GCP, Docker, Kubernetes, CI/CD, GitHub Actions, Git, Containerization",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing",
    },
    "subject": "Application for GenAI Engineer (m/w/d) — Machine Learning Reply",
    "body": [
        "Machine Learning Reply's cross-industry GenAI work for HDAX enterprises is exactly the environment I want my next chapter in — real client problems, LLM-heavy stacks, cloud-native delivery, and end-to-end ownership from strategy to deployment. At IONOS I shipped a production RAG pipeline on Llama 3, my Master's thesis is a working LLM-agent Security Shim (MCP, A2A) with 81% threat-detection recall, and I write Python + LangChain + fine-tuned Llama 3 daily. I'm also B2 in German, English C1 — comfortable in both enterprise worlds.",
        "Here is how my experience maps directly to the GenAI Engineer role:",
        "You build generative AI models, multi-agent LLM applications, and AI pipelines. My thesis is a FastAPI multi-agent orchestrator over MCP and A2A with authentication, policy routing, and measurable evaluation (150 adversarial tests, 98.7% recall). My IONOS RAG pipeline uses Llama 3 on pgvector for enterprise context. I've fine-tuned Llama 3 with QLoRA for DCF reasoning and served it via FastAPI — the end-to-end GenAI loop, not just the notebook.",
        "You want expert Python, prompt/context engineering, fine-tuning, RAG, and cloud deployment (AWS/GCP/Azure). Python is my daily driver; I'm AWS-certified (Bedrock); I've run RAG in production; I've fine-tuned LLMs with Unsloth + QLoRA; and I containerize my ML services with Docker + CI/CD. Databricks and Knowledge Graphs I've worked with academically and am ready to scale in client projects.",
        "You need someone who can present to clients and mentor juniors. At IONOS I presented RAG outcomes directly to senior management; during my Master's I run paper-study sessions; and at KARO I paired with juniors on Spring Boot microservices. I enjoy turning complex AI into something a client stakeholder can actually understand and decide on.",
        "Machine Learning Reply is the kind of consultancy I'd be proud to represent on client sites. I'd bring immediately productive GenAI engineering, a builder mentality shaped across 4 years of cloud and AI delivery, and the rigour I developed researching AI agent security for my thesis. Munich is attractive for me, and I can be on-site quickly.",
    ],
}

# ---------------- 012 GAIA — Senior Product Engineer (LLM) ----------------
JOBS["H2M9"] = {
    "lang": "en",
    "profile": (
        "Product Engineer with 4+ years shipping full-stack software and 2+ years integrating LLM APIs end-to-end. "
        "Strong Python backend, React/Next.js frontend, and direct LLM integration experience (OpenAI-class APIs, "
        "Llama 3, RAG, fine-tuning). Built a production RAG pipeline at IONOS, a FastAPI LLM-agent Security Shim "
        "(MCP, A2A) for my Master's thesis, and a freelance airport parking platform end-to-end in Next.js 16. "
        "Comfortable talking to users, prototyping fast with LLMs, and shipping customer-facing AI features. "
        "CS background, knowledge-intensive domain comfort (financial DCF modelling), German B2, happy to be in Berlin."
    ),
    "thesis_bullets": [
        "Built a FastAPI-based LLM orchestration layer for MCP, A2A, and browser-based AI protocols with policy routing.",
        "Achieved 98.7% recall across 150 adversarial tests; rapid-prototyped prompts and policies for measurable improvements.",
        "Implemented Zero-Trust intent normalisation with EU AI Act-aligned observability for knowledge-intensive AI workflows.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, LLMs, MCP, A2A, Prompt Engineering, Legal/Compliance-class Reasoning, Zero-Trust",
    "ionos_bullets": [
        "Built a production RAG pipeline on PostgreSQL (pgvector) with Llama 3 for internal knowledge-intensive queries.",
        "Rapidly iterated on prompt design with real user feedback; presented outcomes to senior management.",
        "Automated cloud integration tests via the Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, Llama 3, RAG, pgvector, PostgreSQL, FastAPI, LangChain, LLM APIs, Prompt Engineering, CI/CD",
    "genpact_bullets": [
        "Increased user traffic by 30% through an interactive React frontend for enterprise clients.",
        "Ensured 99.9% availability through Docker and CI/CD pipelines.",
        "Engaged directly with enterprise end-users to translate requirements into features.",
    ],
    "genpact_tech": "Tech: React, JavaScript, Docker, CI/CD, REST APIs, Node.js, Git, Customer Research",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, Microservices, MongoDB, Express, React, Node.js, Redux, MERN",
    "skills": {
        "programming": "Programming: Python (Expert), TypeScript, JavaScript, SQL, Java, C",
        "ai_ml": "LLMs & AI: LLM APIs (OpenAI-class, Llama 3), Prompt Design, RAG, Agentic Workflows, LangChain, LlamaIndex, Vector Databases, Fine-Tuning, QLoRA, Hugging Face, PyTorch",
        "llm_infra": "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, Spring AI, LLM Gateway",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI, AI Safety",
        "backend": "Full-Stack: Python (backend), Vue-compatible React / Next.js / React 19, Node.js, REST APIs, OAuth2, Prisma, Stripe",
        "devops": "DevOps / Cloud: AWS (certified), Docker, Kubernetes, CI/CD, GitHub Actions, Vercel, macOS, Git",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing",
    },
    "subject": "Application for Senior Product Engineer (LLM) — GAIA",
    "body": [
        "GAIA caught my attention because it sits at the exact intersection I've been working on: LLMs applied to a knowledge-intensive domain where accuracy, prompt design, and direct user feedback all matter. At IONOS I shipped a production RAG pipeline with Llama 3, my Master's thesis is a FastAPI LLM-agent Security Shim (MCP, A2A) with 81% threat-detection recall, and I've fine-tuned Llama 3 on financial DCF reasoning — the same flavour of compliance-adjacent, structured-reasoning work that legal tech demands.",
        "Here is how my experience maps directly to the Senior Product Engineer role:",
        "You want 2+ years of software engineering across multiple layers, strong Python, and 2-3 LLM projects. I have 4+ years across GenAI, cloud, and full-stack: IONOS (RAG in production), my Master's thesis (LLM-agent orchestration, MCP/A2A), DCF Valuation Engine (fine-tuned Llama 3 served via FastAPI), Dibuco AI (transformer-based sentiment + clustering), and LernKartei AI (Spring AI microservices). That's well past three LLM projects.",
        "You build AI features across Python + Vue/React. Python is my daily driver; I write React 19 and Next.js 16 at production scale (the airport parking platform I just delivered — 22 rate-limited APIs, Stripe, Playwright E2E). Swapping to Vue 3 is idiomatic work for someone already at home in a component-tree mental model.",
        "You want someone who talks to customers and ships. I've done this continuously — presenting RAG results to IONOS senior management, iterating the LIT Airport platform weekly with the ops team, partnering with the KARO founders on deadline-driven features. I'd want to sit close to GAIA's legal-professional users and let their workflow shape the prompts and UI.",
        "GAIA's mission — legal AI that actually helps practitioners — is the kind of work I'd be proud to put my name on. I'd bring immediately productive LLM + full-stack engineering, a builder mentality shaped over 4 years, and the rigour I've picked up researching agent security. Berlin works for me, I'm already based in Germany, and macOS is my daily driver.",
    ],
}

# ---------------- 013 Contentful — Backend Software Engineer ----------------
JOBS["RYRI"] = {
    "lang": "en",
    "profile": (
        "Backend Software Engineer with 4+ years designing and shipping production services in Python, Java, "
        "and TypeScript. Strong system design sense, API versioning instincts, and hands-on with AWS, Docker, "
        "Kubernetes, and CI/CD for services running at high request volume. Built a production RAG pipeline at "
        "IONOS (Llama 3 + PostgreSQL) with cache-aware retrieval, shipped 15+ REST APIs at KARO with 95% test "
        "coverage, and delivered a 22-API Next.js + Prisma platform with observability and rate limiting. "
        "Comfortable owning long-lived production systems and making sound tradeoffs under ambiguity. "
        "German B2, English C1, based in Germany."
    ),
    "thesis_bullets": [
        "Built a FastAPI backend service orchestrating AI protocols (MCP, A2A) with authentication, caching, and rate-limiting.",
        "Achieved 98.7% recall across 150 adversarial tests; designed deterministic policy-engine execution with strong guarantees.",
        "Implemented Zero-Trust intent normalisation with observability over distributed agent components.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, MCP, A2A, Backend Design, Caching, Rate Limiting, Observability",
    "ionos_bullets": [
        "Built a production RAG backend on PostgreSQL (pgvector) with Llama 3 — cache-aware retrieval at request volume.",
        "Presented backend architecture outcomes to senior management; optimised CI/CD for service deployment.",
        "Automated cloud integration tests via Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, PostgreSQL, pgvector, FastAPI, LangChain, Llama 3, AWS, Docker, CI/CD, Distributed Systems",
    "genpact_bullets": [
        "Increased user traffic by 30% through a TypeScript-first frontend wired to REST APIs at high request volume.",
        "Ensured 99.9% availability through Docker containerisation and CI/CD pipelines.",
        "Analysed end-to-end data flows with Celonis Process Mining to debug production system behaviour.",
    ],
    "genpact_tech": "Tech: React, TypeScript, JavaScript, Docker, CI/CD, REST APIs, Node.js, Git, Observability",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, Microservices, API Versioning, MongoDB, Express, React, Node.js, Redux, MERN",
    "skills": {
        "programming": "Programming: TypeScript, Python (Expert), Java, JavaScript, SQL, C",
        "ai_ml": "AI / ML: LLMs, RAG, Agentic Workflows, LangChain, Vector Databases, Fine-Tuning, QLoRA, PyTorch, Hugging Face",
        "llm_infra": "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, Spring AI, LLM Gateway",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI",
        "backend": "Backend: Spring Boot, Microservices, REST APIs, GraphQL-ready, API Versioning, Schemas, Serverless, Distributed Systems, Data Transformation Pipelines, Caching, Prisma, Node.js, Express",
        "devops": "DevOps / Cloud: AWS (certified), CloudFlare, Docker, Kubernetes, Terraform, CI/CD, GitHub Actions, Git, Observability",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing, Observability (telemetry, logging, tracing)",
    },
    "subject": "Application for Backend Software Engineer (f/m/d) — Contentful",
    "body": [
        "Contentful's Entity Architecture team is working on the kind of backend problem I enjoy most: long-lived services, API + schema evolution at billions of requests per month, and cache-aware execution with real performance guarantees. At IONOS I shipped a production RAG pipeline on PostgreSQL (pgvector) with Llama 3; at KARO I delivered 15+ REST APIs with 95% test coverage; and for a freelance airport client I shipped 22 rate-limited APIs with JWT auth and observability on Vercel + Neon. I like the intersection of content platforms and AI tooling Contentful is positioned at.",
        "Here is how my experience maps directly to the Backend Software Engineer role:",
        "You design and build core backend systems with API/schema evolution and strong performance guarantees. At IONOS I designed a cache-aware RAG pipeline; at KARO I evolved 15+ Spring Boot REST APIs while keeping 95% test coverage. In my freelance work I own API versioning, JWT auth, rate limiting, and observability across 22 endpoints — deterministic, debuggable, and safe to evolve.",
        "You run on AWS, CloudFlare, Docker, Kubernetes, and Terraform. I'm AWS-certified, ship services in Docker at 99.9% availability, pair Kubernetes with CI/CD for ML model deployment at IONOS, and am comfortable with Infrastructure-as-Code mental models (Terraform-style declarative design). CloudFlare fits the caching instincts I already apply.",
        "You value TypeScript, GraphQL, and AI tooling. My recent freelance work is TypeScript + Next.js 16 + Prisma; my day-to-day AI stack is LangChain + FastAPI + pgvector; and I've fine-tuned Llama 3 with QLoRA. The declarative-systems/data-transformation/cache-heavy architecture space is where I feel at home.",
        "Contentful is exactly the kind of team where I'd want to own production systems for years. I'd bring immediately productive backend engineering, system-design judgement shaped over 4 years of Spring Boot, FastAPI, and Next.js, and the rigour I developed researching AI agent security for my thesis. Berlin works for me, and I'm comfortable in ambiguity.",
    ],
}

# ---------------- 014 Contentful — Fullstack Software Engineer ----------------
JOBS["MKWH"] = {
    "lang": "en",
    "profile": (
        "Fullstack Software Engineer with 4+ years shipping TypeScript, React, Node.js, and Python across "
        "backend and frontend. Comfortable with Docker, Kubernetes, CI/CD, and cloud platforms (AWS certified). "
        "Built a production RAG pipeline at IONOS (Llama 3 + PostgreSQL), delivered a Next.js 16 + React 19 + "
        "Prisma airport parking platform end-to-end for a freelance client (22 APIs, Stripe, Playwright E2E), "
        "and shipped a React frontend at GENPACT that lifted traffic 30% with 99.9% availability. Strong "
        "problem-solving, product-first mindset, and used to juggling competing priorities. German B2, English C1, "
        "based in Germany and happy to be in Berlin."
    ),
    "thesis_bullets": [
        "Built a FastAPI backend + lightweight dashboard orchestrating AI protocols (MCP, A2A) with authentication and rate-limiting.",
        "Achieved 98.7% recall across 150 adversarial tests; implemented end-to-end workflow telemetry for collaboration.",
        "Implemented Zero-Trust intent normalisation with EU AI Act-aligned observability across distributed components.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, TypeScript, MCP, A2A, Observability, Distributed Systems",
    "ionos_bullets": [
        "Built a production RAG pipeline on PostgreSQL (pgvector) with Llama 3 — full-stack flows from API to UI.",
        "Presented RAG outcomes to senior management; optimised CI/CD for ML services.",
        "Automated cloud integration tests via Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, PostgreSQL, pgvector, FastAPI, LangChain, Llama 3, AWS, Docker, CI/CD, TypeScript, Node.js",
    "genpact_bullets": [
        "Increased user traffic by 30% through an interactive, responsive React + TypeScript frontend.",
        "Ensured 99.9% availability through Docker containerisation and CI/CD pipelines.",
        "Analysed end-to-end data flows with Celonis Process Mining to inform product iterations.",
    ],
    "genpact_tech": "Tech: React, TypeScript, JavaScript, Docker, CI/CD, REST APIs, Node.js, Git",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, MongoDB, Express, React, Node.js, Redux, MERN, TypeScript",
    "skills": {
        "programming": "Programming: TypeScript, Python (Expert), JavaScript, Java, SQL, C",
        "ai_ml": "AI / ML: LLMs, RAG, LangChain, Vector Databases, Fine-Tuning, QLoRA, PyTorch, Hugging Face, AI Tooling Integration",
        "llm_infra": "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, Spring AI",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI",
        "backend": "Full-Stack: TypeScript, React 19, Node.js, Next.js 16, Express, REST APIs, GraphQL-ready, OAuth2, JWT, Prisma, Stripe, Tailwind, WebSockets",
        "devops": "DevOps / Cloud: AWS (certified), Docker, Kubernetes, CI/CD, GitHub Actions, Vercel, Neon, Git, Observability",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing",
    },
    "subject": "Application for Fullstack Software Engineer (f/m/d) — Contentful",
    "body": [
        "Contentful's Automations, Workflows, and Collaboration Tools team is the kind of fullstack problem I love: features that thousands of customers rely on, in a TypeScript + React + Node.js stack with Kubernetes underneath. I've spent the last 4 years shipping exactly that combination. At IONOS I shipped a production RAG pipeline with Llama 3; my recent freelance work is Next.js 16 + React 19 + Prisma + 22 REST APIs for a real airport client; and at GENPACT I drove 30% traffic with 99.9% availability.",
        "Here is how my experience maps directly to the Fullstack Software Engineer role:",
        "You need strong TypeScript, React, and Node.js, with Docker and Kubernetes in the stack. TypeScript is my daily driver for the airport platform — 22 rate-limited APIs, Playwright E2E, Stripe, Twilio SMS, JWT auth, Next.js 16 + React 19. At IONOS and GENPACT I containerised services with Docker and paired CI/CD with Kubernetes-class deployment. I think about the full stack as one system.",
        "You want cloud knowledge (AWS/Azure/GCP) and product-oriented fullstack engineering. I'm AWS-certified (Bedrock); I ship to Vercel + Neon + GitHub Actions for the airport platform; I built a React frontend at GENPACT that lifted traffic 30%; and I've run a production RAG pipeline at IONOS. I care about product metrics, not just code shipping.",
        "You work on automations, workflows, and collaboration features. My Master's thesis is essentially a workflow-orchestration engine for AI agents (MCP, A2A, policy-routed collaboration between components). The airport platform is a workflow-heavy booking + admin dashboard with rewards, rate limits, and multi-role flows. The mental model transfers cleanly.",
        "Contentful is the kind of team I'd want to do my best fullstack work on — meaningful scale, modern TS/Node/Kubernetes stack, and real customer impact. I'd bring immediately productive fullstack + AI tooling experience, a builder mentality shaped across 4 years of shipping, and the discipline that comes from researching AI agent security in my thesis.",
    ],
}

# ---------------- 015 SumUp — Senior AI/ML Engineer ----------------
JOBS["JRST"] = {
    "lang": "en",
    "profile": (
        "Senior-track AI/ML Engineer with 4+ years of Python production experience, shipping LLM-powered "
        "services, RAG pipelines, agentic workflows, and fine-tuned models. Hands-on with FastAPI, LangChain, "
        "vector databases (pgvector), Docker, Kubernetes, CI/CD, and MLOps tooling (MLflow, Airflow, Langfuse). "
        "Built a production RAG pipeline at IONOS (Llama 3 + PostgreSQL), a Security Shim for agentic protocols "
        "(MCP, A2A) with 81% threat-detection recall for my Master's thesis, and a fine-tuned Llama 3 DCF engine "
        "served via FastAPI. Comfortable with LLM evaluation, production model deployment, and cross-functional "
        "collaboration. AWS-certified, German B2, English C1, happy to be office-first in Berlin."
    ),
    "thesis_bullets": [
        "Built a FastAPI-based Security Shim / LLM orchestration layer for MCP, A2A, and browser-based AI agent protocols with LLM evaluation hooks.",
        "Achieved 98.7% recall across 150 adversarial tests; production-grade LLM evaluation with zero unauthorised tool calls.",
        "Implemented Zero-Trust intent normalisation with telemetry, logging, and tracing over distributed agent components.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, MCP, A2A, LLM Evaluation, MLOps, Zero-Trust, Observability",
    "ionos_bullets": [
        "Built a production RAG pipeline on PostgreSQL (pgvector) with Llama 3 — end-to-end MLOps from dataset prep to serving.",
        "Presented RAG outcomes to senior management; optimised CI/CD for ML model deployment and monitoring.",
        "Automated cloud integration tests via Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, Llama 3, RAG, pgvector, PostgreSQL, FastAPI, LangChain, AWS, Docker, Kubernetes, CI/CD, MLOps, Airflow-style orchestration",
    "genpact_bullets": [
        "Increased user traffic by 30% through a production-grade React frontend backed by REST APIs.",
        "Ensured 99.9% availability through Docker containerisation and CI/CD pipelines.",
        "Analysed end-to-end data flows with Celonis Process Mining to identify bottlenecks.",
    ],
    "genpact_tech": "Tech: React, JavaScript, Docker, CI/CD, REST APIs, Node.js, Git, Monitoring",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, Microservices, MongoDB, Express, React, Node.js, Redux, MERN",
    "skills": {
        "programming": "Programming: Python (Expert), Java, TypeScript, JavaScript, SQL, C",
        "ai_ml": "AI / ML: LLMs, RAG, Agentic Workflows, LLM Evaluation, Chatbot Assistants, AI Translation, LLM Orchestration, LangChain, LlamaIndex, Vector Databases, Fine-Tuning, QLoRA, PEFT, PyTorch, TensorFlow, Hugging Face, Prompt Engineering, Deep Learning",
        "llm_infra": "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, GenAI/LLM APIs, Large Text Datasets",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI, AI Safety",
        "backend": "Backend: Spring Boot, Microservices, REST APIs, Data Pipelines, Distributed Systems, API Architecture",
        "devops": "Cloud / MLOps: AWS (certified, Bedrock, S3), GCP, Azure, Docker, Kubernetes, CI/CD, GitHub Actions, MLflow, Kubeflow, Airflow, Langfuse, Prometheus, CloudWatch, Git",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing, Model Monitoring, Alerting",
    },
    "subject": "Application for Senior AI/ML Engineer — SumUp",
    "body": [
        "SumUp's Operations AI team — an AI Assistant already automating ~40% of merchant requests, plus an AI Translation layer and an Agent Copilot — is exactly the class of problem I've been building against. At IONOS I shipped a production RAG pipeline on Llama 3, my Master's thesis is a FastAPI-based LLM-orchestration Security Shim over MCP and A2A with 81% threat-detection recall, and I fine-tune LLMs with QLoRA and serve them via FastAPI. The muscle SumUp needs — production ML/AI, LLM evaluation, MLOps — is the muscle I've been training.",
        "Here is how my experience maps directly to the Senior AI/ML Engineer role:",
        "You design, deploy, and maintain core AI products: Assistant, Translation, Agent Copilot. At IONOS I designed a RAG pipeline with Llama 3 + pgvector; my thesis orchestrates agent protocols (MCP, A2A) with authentication, policy routing, and measurable evaluation (150 adversarial tests). I've fine-tuned Llama 3 on DCF reasoning and served inference via FastAPI — end-to-end LLM product work, not just notebooks.",
        "You need expert Python, production ML deep learning, and MLOps tooling (MLflow, Kubeflow, Airflow, Langfuse). Python is my daily driver; I've shipped containerised ML services at 99.9% availability; I'm AWS-certified (Bedrock); and I care about observability (telemetry, logging, tracing) in every system I build. Kubernetes + Docker + CI/CD is how my services reach production.",
        "You need someone who prepares large text datasets and evaluates LLMs rigorously. My thesis is literally an adversarial dataset + LLM evaluation loop — 150 tests, 98.7% recall, policy-level ablations. At IONOS I shaped support data for RAG retrieval; Dibuco AI is an NYT data pipeline with K-Means + transformer sentiment. I'm comfortable at both the data-prep and the evaluation ends.",
        "SumUp Operations AI is the kind of product team I'd want to sit with for years — real merchants, real latency constraints, real product metrics. I'd bring immediately productive AI engineering, a builder mentality shaped over 4 years of cloud and GenAI, and the rigour I picked up researching AI agent security. Berlin office-first works for me — I'm already based in Germany.",
    ],
}

# ---------------- 016 SumUp — Senior ML Engineer (Sofia) ----------------
JOBS["LA6K"] = {
    "lang": "en",
    "profile": (
        "Senior-track Machine Learning Engineer with 4+ years of Python production experience and 2+ years "
        "hands-on with LLMs, RAG systems, and chatbot-style conversational AI. Shipped a production RAG "
        "pipeline at IONOS (Llama 3 + PostgreSQL), a FastAPI LLM-agent Security Shim (MCP, A2A) with 81% "
        "threat-detection recall for my Master's thesis, and a fine-tuned Llama 3 DCF reasoning engine. "
        "Comfortable with full ML workflow (feature engineering to serving + monitoring + alerting), MLOps "
        "tooling (MLflow, Airflow, Langfuse), and cloud deployment on AWS. German B2, English C1, open to "
        "Sofia for the right team."
    ),
    "thesis_bullets": [
        "Built a FastAPI LLM-agent Security Shim for MCP, A2A, and browser-based AI protocols — chatbot-adjacent orchestration.",
        "Achieved 98.7% recall across 150 adversarial tests; implemented full ML workflow monitoring and alerting.",
        "Implemented Zero-Trust intent normalisation with observability over distributed agent components.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, MCP, A2A, LLM Evaluation, Monitoring, Alerting, MLOps",
    "ionos_bullets": [
        "Built a production RAG pipeline on PostgreSQL (pgvector) with Llama 3 — full ML workflow from feature engineering to serving.",
        "Presented RAG outcomes to senior management; optimised CI/CD for ML deployment and monitoring.",
        "Automated cloud integration tests via Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, Llama 3, RAG, pgvector, PostgreSQL, FastAPI, LangChain, AWS, Docker, MLOps, CI/CD",
    "genpact_bullets": [
        "Increased user traffic by 30% through a responsive React frontend backed by REST APIs.",
        "Ensured 99.9% availability through Docker containerisation and CI/CD pipelines.",
        "Analysed end-to-end data flows with Celonis Process Mining to identify bottlenecks.",
    ],
    "genpact_tech": "Tech: React, JavaScript, Docker, CI/CD, REST APIs, Node.js, Git, Monitoring",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, Microservices, MongoDB, Express, React, Node.js, Redux, MERN",
    "skills": {
        "programming": "Programming: Python (Expert), Java, TypeScript, JavaScript, SQL, C",
        "ai_ml": "AI / ML: LLMs, RAG Systems, Conversational AI, Chatbot Assistants, Agentic Workflows, LangChain, LlamaIndex, Vector Databases, Fine-Tuning, QLoRA, PEFT, PyTorch, TensorFlow, Hugging Face, Prompt Engineering, Deep Learning",
        "llm_infra": "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, GenAI APIs, Large Text Datasets",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI, AI Safety",
        "backend": "Backend: Spring Boot, Microservices, REST APIs, Data Pipelines, API Integration",
        "devops": "Cloud / MLOps: AWS (certified), GCP, Azure, Docker, Kubernetes, CI/CD, GitHub Actions, MLflow, Kubeflow, Airflow, Langfuse, Git",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing, Model Monitoring, Alerting",
    },
    "subject": "Application for Senior Machine Learning Engineer — SumUp",
    "body": [
        "SumUp's Global Operations AI team — models and systems powering intelligent customer support for 4M+ merchants — is the exact class of ML engineering I want to be doing: RAG systems, chatbot assistants, and the full ML workflow from data prep through serving and monitoring. At IONOS I shipped a production RAG pipeline on Llama 3, my Master's thesis is a working FastAPI LLM-agent Security Shim (MCP, A2A) with 81% threat-detection recall, and I fine-tune LLMs with QLoRA for reasoning-heavy tasks.",
        "Here is how my experience maps directly to the Senior Machine Learning Engineer role:",
        "You build and deploy scalable ML/AI in production with RAG and chatbot-style systems. My IONOS RAG pipeline is exactly that: Llama 3 + pgvector serving production queries. My thesis adds agent orchestration (MCP, A2A) on top, with measurable LLM evaluation (150 adversarial tests, 98.7% recall). I've also fine-tuned Llama 3 with Unsloth + QLoRA and served inference via FastAPI.",
        "You need advanced Python + MLOps tools (MLflow, Kubeflow, Airflow, Langfuse). Python is my daily driver; I containerize ML services with Docker and ship via CI/CD; I pair Kubernetes-class deployment with observability; and I'm AWS-certified (Bedrock). MLflow + Airflow + Langfuse slot naturally into the way I already structure workflows.",
        "You own the full ML workflow from feature engineering through serving, monitoring, and alerting. At IONOS I shaped support-ticket data for RAG retrieval and cared about the serving path end-to-end. In my thesis I built telemetry, logging, and tracing into every agent run. The monitoring + alerting mindset is baked in.",
        "SumUp's AI team looks like the kind of place where I'd do my best applied-ML work — real merchants, latency constraints, product impact. I'd bring immediately productive ML engineering, a builder mentality shaped over 4 years of cloud and GenAI, and the rigour I picked up researching AI agent security for my thesis. Open to Sofia for the right mission.",
    ],
}

# ---------------- 017 SumUp — Senior ML Engineer (Barcelona) ----------------
JOBS["A99F"] = {
    "lang": "en",
    "profile": (
        "Senior-track Machine Learning Engineer with 4+ years of Python production experience and hands-on "
        "conversational AI and RAG system delivery. Built a production RAG pipeline at IONOS (Llama 3 + "
        "PostgreSQL), a FastAPI LLM-agent Security Shim (MCP, A2A) with 81% threat-detection recall for my "
        "Master's thesis, and a fine-tuned Llama 3 engine served via FastAPI. Comfortable across the full ML "
        "workflow — feature engineering, fine-tuning, serving, monitoring, alerting — with MLOps tooling "
        "(MLflow, Airflow, Langfuse) and cloud deployment on AWS. AWS-certified, German B2, English C1, open "
        "to Barcelona."
    ),
    "thesis_bullets": [
        "Built a FastAPI LLM-agent Security Shim for MCP, A2A, and browser-based AI protocols — conversational AI-adjacent orchestration.",
        "Achieved 98.7% recall across 150 adversarial tests; implemented full ML workflow monitoring and alerting.",
        "Implemented Zero-Trust intent normalisation with observability over distributed agent components.",
    ],
    "thesis_tech": "Tech: FastAPI, Python, MCP, A2A, LLM Evaluation, Monitoring, Alerting, MLOps",
    "ionos_bullets": [
        "Built a production RAG pipeline on PostgreSQL (pgvector) with Llama 3 — full ML workflow from feature engineering to serving at scale.",
        "Presented RAG outcomes to senior management; optimised CI/CD for ML deployment and monitoring.",
        "Automated cloud integration tests via Karate Framework (Java) — 30% reduction in manual QA effort.",
    ],
    "ionos_tech": "Tech: Python, Llama 3, RAG, pgvector, PostgreSQL, FastAPI, LangChain, AWS, Docker, MLOps, CI/CD",
    "genpact_bullets": [
        "Increased user traffic by 30% through a responsive React frontend at large-scale production volume.",
        "Ensured 99.9% availability through Docker containerisation and CI/CD pipelines.",
        "Analysed end-to-end data flows with Celonis Process Mining to identify bottlenecks.",
    ],
    "genpact_tech": "Tech: React, JavaScript, Docker, CI/CD, REST APIs, Node.js, Git",
    "iit_tech": "Tech: Java, Spring Boot, JUnit, REST APIs, Microservices, MongoDB, Express, React, Node.js, Redux, MERN",
    "skills": {
        "programming": "Programming: Python (Expert), Java, TypeScript, JavaScript, SQL, C",
        "ai_ml": "AI / ML: LLMs, RAG Systems, Conversational AI, Chatbot Assistants, Agentic Workflows, LangChain, LlamaIndex, Vector Databases, Fine-Tuning, QLoRA, PEFT, PyTorch, TensorFlow, Hugging Face, Prompt Engineering",
        "llm_infra": "LLM Infra: AWS Bedrock, PostgreSQL, pgvector, FastAPI, Llama 3, GenAI APIs, Large Text Datasets",
        "ai_gov": "AI Governance: Guardrails, Zero-Trust, MCP, A2A, EU AI Act, Responsible AI",
        "backend": "Backend: Spring Boot, Microservices, REST APIs, Data Pipelines, API Integration",
        "devops": "Cloud / MLOps: AWS (certified), GCP, Azure, Docker, Kubernetes, CI/CD, GitHub Actions, MLflow, Kubeflow, Airflow, Langfuse, Git",
        "testing": "Testing: JUnit, Karate, pytest, Playwright, E2E, Unit & Integration Testing, Model Monitoring, Alerting",
    },
    "subject": "Application for Senior Machine Learning Engineer — SumUp",
    "body": [
        "SumUp's Global Operations AI team — 4M+ merchants served by intelligent support systems — is the exact kind of applied-ML work I want to do next: RAG, conversational AI, chatbot assistants, and the full ML workflow end-to-end. At IONOS I shipped a production RAG pipeline on Llama 3, my Master's thesis is a FastAPI LLM-agent Security Shim (MCP, A2A) with 81% threat-detection recall, and I fine-tune LLMs with QLoRA for reasoning-heavy tasks.",
        "Here is how my experience maps directly to the Senior Machine Learning Engineer role:",
        "You architect, build, and deploy AI in large-scale production with RAG and conversational AI. My IONOS RAG pipeline is exactly that class — Llama 3 + pgvector serving production queries at enterprise load. My thesis adds agent orchestration (MCP, A2A) with measurable LLM evaluation (150 adversarial tests, 98.7% recall). I've also fine-tuned Llama 3 with Unsloth + QLoRA and served inference via FastAPI.",
        "You need advanced Python + MLOps tools (MLflow, Kubeflow, Airflow, Langfuse). Python is my daily driver; I containerize ML services with Docker and deploy via CI/CD; I pair Kubernetes-class deployment with observability; I'm AWS-certified (Bedrock). The MLflow + Airflow + Langfuse stack maps naturally onto how I already structure workflows.",
        "You own the full ML workflow from feature engineering through serving, monitoring, and alerting. At IONOS I shaped support-ticket data for RAG retrieval; in my thesis I built telemetry, logging, and tracing into every agent run. Collecting, preprocessing, and preparing large text datasets is as much my daily work as model training.",
        "SumUp Barcelona looks like a team where I'd do my best applied-ML work — real merchants, real latency budgets, real product metrics. I'd bring immediately productive ML engineering, a builder mentality shaped over 4 years of cloud and GenAI, and the rigour that comes from researching AI agent security. Open to Barcelona for this role.",
    ],
}


# ============================================================================

def run_one(ref):
    cfg = JOBS[ref]
    if cfg["lang"] == "en":
        cv_out = tailor_cv_en(ref, cfg)
        cl_out = tailor_cl_en(ref, cfg)
    else:
        raise ValueError("DE lang not implemented in this batch")
    print(f"DONE {ref}: {cv_out}, {cl_out}")
    return cv_out, cl_out


if __name__ == "__main__":
    ref = sys.argv[1]
    if ref == "all":
        for r in JOBS:
            try:
                run_one(r)
            except Exception as e:
                print(f"FAILED {r}: {e}", file=sys.stderr)
    else:
        run_one(ref)
